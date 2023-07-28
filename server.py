from flask import Flask, request, render_template, jsonify, make_response, send_file, url_for, redirect, session
from flask_socketio import SocketIO

import openai
import os

import pdfplumber
import nltk
nltk.download('punkt')
nltk.download('stopwords')
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import string
from wordcloud import WordCloud
from docx import Document
import io
from azure.storage.fileshare import ShareClient, ShareFileClient
from azure.storage.blob import ContainerClient, BlobServiceClient, BlobClient
import tempfile

import logging
"""logging.basicConfig(level=logging.INFO)

handler = logging.StreamHandler()
handler.setLevel(logging.INFO)"""

job_desc_text = ""
resume_text = ""

job_desc_url = ""
resume_url = ""

resume_filename = ""
job_desc_filename = ""

app = Flask(__name__)
app.secret_key = "INSERT_APP_SECRET_KEY_HERE"
openai.api_key = "INSERT_API_KEY_HERE"

socketio = SocketIO(app)

def save_file_to_azure(file, file_name):
    sas_url = "https://resumeverifierstorage.blob.core.windows.net/files?sp=racwdli&st=2023-06-23T19:33:11Z&se=2025-07-24T03:33:11Z&sv=2022-11-02&sr=c&sig=ReLjjFuQE174X8C3ba0%2F4yk6gXfLw0%2F1tnjQ9z18S7Q%3D"
    container_client = ContainerClient.from_container_url(sas_url)
    blob_client = container_client.get_blob_client(file_name)

    # Check if the file already exists
    if blob_client.exists():
        # If the file exists, delete it before uploading the new file
        blob_client.delete_blob()

    temp_dir = tempfile.TemporaryDirectory()
    temp_file_path = os.path.join(temp_dir.name, file_name)
    file.save(temp_file_path)

    with open(temp_file_path, "rb") as data:
        blob_client.upload_blob(data)
    
    temp_dir.cleanup()
    blob_url = blob_client.url
    print(blob_url)

    return blob_url

def open_file_from_azure(file_path):
    storage_connection_string = os.environ["StorageConnection"]
    share_name = "resumeverifierfiles"

    with ShareFileClient.from_connection_string(storage_connection_string, share_name, file_path) as file_client:
        file_content = file_client.download_file().readall()

    return file_content

@app.route('/render-uploaded-files')
def render_uploaded_files():
    global resume_url
    global job_desc_url
    print(resume_url)
    print(job_desc_url)
    print('im here!!!!')
    
    return render_template('uploaded_files.html', resume_url=resume_url, job_desc_url=job_desc_url)

@app.route('/delete-files', methods=['POST'])
def delete_files():
    global resume_filename
    global job_desc_filename

    sas_url = "https://resumeverifierstorage.blob.core.windows.net/files?sp=racwdli&st=2023-06-23T19:33:11Z&se=2025-07-24T03:33:11Z&sv=2022-11-02&sr=c&sig=ReLjjFuQE174X8C3ba0%2F4yk6gXfLw0%2F1tnjQ9z18S7Q%3D"
    container = ContainerClient.from_container_url(sas_url)
    blobs_list = container.list_blobs()
    print(resume_filename)
    print(job_desc_filename)
    for blob in blobs_list:
        if blob.name == resume_filename or blob.name == job_desc_filename:
            print(blob.name)
            container.delete_blob(blob)
    
    return jsonify({"message": "Files deleted"}), 200

@app.route('/remove-images', methods=['POST'])
def remove_images():
    global job_desc_text
    global resume_text
    job_desc_text=""
    resume_text=""
    return "deleted"

# Function to preprocess text for creating word clouds
stop_words = set(stopwords.words('english'))
def preprocess_text(text):
    # Convert to lower case
    text = text.lower()
    # Tokenize
    word_tokens = word_tokenize(text)
    # Remove stop words and punctuation
    preprocessed_text = [word for word in word_tokens if word not in stop_words and word not in string.punctuation]
    # Remove specific words
    words_to_exclude = ['project', 'candidate' 'requirement','required','must','first','previously','experience','provided','using','years']  # Add the words you want to exclude
    preprocessed_text = [word for word in preprocessed_text if word not in words_to_exclude]
    return ' '.join(preprocessed_text)

@app.route('/api/get_resume_wordcloud')
def get_resume_wordcloud():
    global resume_text
    resume_skills = preprocess_text(job_desc_text)
    resume_wordcloud = WordCloud(width=1000, height=500).generate(resume_skills)

    # Convert the word cloud image to bytes
    image_bytes = io.BytesIO()
    resume_wordcloud.to_image().save(image_bytes, format='PNG')
    image_bytes.seek(0)

    # Return the image as a response
    return send_file(image_bytes, mimetype='image/png')

@app.route('/api/get_jobdesc_wordcloud')
def get_jobdesc_wordcloud():
    global job_desc_text
    jobdesc_skills = preprocess_text(job_desc_text)
    jobdesc_wordcloud = WordCloud(width=1000, height=500).generate(jobdesc_skills)

    # Convert the word cloud image to bytes
    image_bytes = io.BytesIO()
    jobdesc_wordcloud.to_image().save(image_bytes, format='PNG')
    image_bytes.seek(0)

    # Return the image as a response
    return send_file(image_bytes, mimetype='image/png')

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/download-results', methods=['POST'])
def download_results():
    result = request.form.get('result')
    response = make_response(result)
    response.headers.set('Content-Disposition', 'attachment', filename='result.txt')
    response.headers.set('Content-Type', 'text/plain')
    return response

@app.route('/upload', methods=['POST'])
def upload_files():
    global job_desc_text
    global resume_text
    global job_desc_url
    global resume_url
    global resume_filename
    global job_desc_filename
    job_desc_text = ""
    resume_text = ""

    resume = request.files.get('resume')
    job_desc = request.files.get('job_desc')

    resume_filename = resume.filename
    job_desc_filename = job_desc.filename

    resume_extension = resume_filename.rsplit('.', 1)[1].lower()
    job_desc_extension = job_desc_filename.rsplit('.', 1)[1].lower()

    if resume_extension == 'pdf':
        print('resume is pdf\n')

        displayer = "https://docs.google.com/viewer?url="
        ending = "&embedded=true"

        resume_url = save_file_to_azure(resume, resume_filename)
        resume_url = displayer+resume_url+ending

        # Convert the pdf files to text
        with pdfplumber.open(resume) as pdf:
            resume_text = '\n'.join(page.extract_text() for page in pdf.pages)
        print(resume_text+'\n')

    elif resume_extension in ['doc', 'docx']:
        print('resume is doc\n')

        displayer = "https://docs.google.com/viewer?url="
        ending = "&embedded=true"

        resume_url = save_file_to_azure(resume, resume_filename)
        resume_url = displayer+resume_url+ending

        doc = Document(resume)

        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text not in paragraphs:
                paragraphs.append(paragraph.text)

        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    #print(cell.text)
                    if cell.text not in tables_text:   
                        tables_text.append(cell.text)

        resume_text = '\n'.join(paragraphs + tables_text)
        print(resume_text+'\n')  # Check the final resume text

    if job_desc_extension == 'pdf':
        print('job desc is a pdf\n')

        displayer = "https://docs.google.com/viewer?url="
        ending = "&embedded=true"
        
        job_desc_url = save_file_to_azure(job_desc, job_desc_filename)
        job_desc_url = displayer+job_desc_url+ending

        # Convert the pdf files to text
        with pdfplumber.open(job_desc) as pdf:
            job_desc_text = '\n'.join(page.extract_text() for page in pdf.pages)
        
        print(job_desc_text+'\n')

    elif job_desc_extension in ['doc', 'docx']:
        print('job desc is a doc\n')

        displayer = "https://docs.google.com/viewer?url="
        ending = "&embedded=true"

        job_desc_url = save_file_to_azure(job_desc, job_desc_filename)
        job_desc_url = displayer+job_desc_url+ending

        doc = Document(job_desc)

        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text not in paragraphs:
                paragraphs.append(paragraph.text)

        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text not in tables_text:   
                        tables_text.append(cell.text)

        job_desc_text = '\n'.join(paragraphs + tables_text)
        print(job_desc_text+'\n')  # Check the final resume text
    
    prompt = request.form.get('prompt')

    # Compare the resume to the job description using the OpenAI API.
    
    try:
        print('trying model \n')
        #print(resume_text + '\n')
        #print(job_desc_text + '\n')
        response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo-16k",
        messages=[
            {"role": "system", "content": "You are an assistant to a recruiter that helps compare resumes against job descriptions. The user will provide you with a resume and a job description."},
            {"role": "user", "content": f"Resume: {resume_text}\n ====== \n Job Description: {job_desc_text}\n ====== \n {prompt}"}
        ]
        )
        #print('done w api call!!')
        #print(response)
        #Find the last message from the assistant
        assistant_message = response['choices'][0]['message']['content']
        print('assistant message: \n')
        print(assistant_message)
    
    except Exception as e:
        # In case of API error, return the error message
        e = str(e)
        output = "There was an error! Here is the message: " + e
        #print(output)
        return jsonify(output)

    return jsonify(assistant_message)

if __name__ == "__main__":
    app.run(debug=True)
